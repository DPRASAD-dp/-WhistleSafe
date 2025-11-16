# submission_verification.py
import os
import atexit
from pathlib import Path

import streamlit as st
import sqlite3
from docx import Document

from dfpipeline import analyze_video
from gemini_processing import process_question_with_doc, setup_gemini
from lingo_translation import translate, LANGUAGES

# -------------------------
# Sidebar: Language selector
# -------------------------
st.sidebar.header("🌐 " + translate("Select Language", "en"))  # show default english label in sidebar header
selected_language = st.sidebar.selectbox(
    translate("Language", "en"),
    list(LANGUAGES.keys()),
    index=list(LANGUAGES.keys()).index("English") if "English" in LANGUAGES else 0,
    key="language_select"
)
TARGET_LANG = LANGUAGES[selected_language]

# -------------------------
# Custom CSS (kept as-is but text values translated in UI)
# -------------------------
st.markdown(
    """
    <style>
    .stApp {
        background: linear-gradient(to bottom right, #1e3c72, #2a5298);
    }
    .css-1d391kg, .css-12oz5g7 {
        background-color: rgba(255, 255, 255, 0.9);
        padding: 20px;
        border-radius: 10px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .stButton button {
        background-color: #2e5cb8;
        color: white;
        border: none;
        border-radius: 5px;
    }
    .stButton button:hover {
        background-color: #3d7be3;
    }
    h1, h2, h3 {
        color: white !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# -------------------------
# Database connection (cached)
# -------------------------
@st.cache_resource
def get_database_connection():
    conn = sqlite3.connect("user_data.db", check_same_thread=False)
    return conn


conn = get_database_connection()
cursor = conn.cursor()


def get_all_reports():
    cursor.execute("SELECT id, user_id, video_path, text_report, status FROM uploads")
    return cursor.fetchall()


def update_status(report_id, status):
    cursor.execute("UPDATE uploads SET status = ? WHERE id = ?", (status, report_id))
    conn.commit()


def verify_status(report_id):
    cursor.execute("SELECT status FROM uploads WHERE id = ?", (report_id,))
    row = cursor.fetchone()
    return row[0] if row else None


# -------------------------
# Report docx generation
# -------------------------
def generate_report(report_id, video_analysis_results, text_analysis_results_en, status_en):
    """
    - video_analysis_results: leave in English (as requested)
    - text_analysis_results_en: Gemini output in English
    - status_en: status string in English ('Accepted'/'Rejected')
    We will translate Gemini output and status for the generated doc.
    """
    reports_folder = Path("reports")
    reports_folder.mkdir(parents=True, exist_ok=True)
    file_path = reports_folder / f"report_{report_id}.docx"

    doc = Document()
    doc.add_heading(f"Report ID: {report_id}", level=1)

    # Video results (keep in English)
    if video_analysis_results:
        doc.add_heading("Video Analysis Results:", level=2)
        doc.add_paragraph(str(video_analysis_results))

    # Text analysis: translate Gemini output into TARGET_LANG
    if text_analysis_results_en:
        doc.add_heading("Text Analysis Results:", level=2)
        translated_doc_text = translate(str(text_analysis_results_en), TARGET_LANG)
        doc.add_paragraph(translated_doc_text)

    doc.add_heading("Final Status:", level=2)
    translated_status = translate(status_en, TARGET_LANG)
    doc.add_paragraph(translated_status)

    doc.save(str(file_path))
    return str(file_path)


# -------------------------
# Session state initialization helper
# -------------------------
def init_session_state(report_id):
    video_key = f"video_analysis_{report_id}"
    text_key = f"text_analysis_en_{report_id}"  # store Gemini output in English in state
    if video_key not in st.session_state:
        st.session_state[video_key] = None
    if text_key not in st.session_state:
        st.session_state[text_key] = None


# Ensure reports folder exists
Path("reports").mkdir(exist_ok=True)

# -------------------------
# Page title
# -------------------------
st.title(translate("WhistleSafe : Authority Dashboard", TARGET_LANG))

# -------------------------
# Initialize Gemini LLM
# -------------------------
gemini_llm, gemini_msg = setup_gemini()
if not gemini_llm:
    st.error(translate("Failed to initialize Gemini. Please check API configuration.", TARGET_LANG))
else:
    st.info(translate(gemini_msg, TARGET_LANG))

# -------------------------
# Fetch and display reports
# -------------------------
reports = get_all_reports()

if not reports:
    st.info(translate("No reports available for review.", TARGET_LANG))
else:
    # Iterate reports
    for report in reports:
        report_id, user_id, video_path, text_report_en, current_status_en = report
        init_session_state(report_id)

        # Container per report
        with st.container():
            st.markdown("### " + translate(f"Report ID: {report_id}", TARGET_LANG))
            st.write(translate(f"User ID: {user_id}", TARGET_LANG))

            # Video display
            if video_path and os.path.exists(video_path):
                st.video(video_path)
            else:
                st.warning(translate(f"Video file not found: {video_path}", TARGET_LANG))

            # Display text report (translated for UI)
            # We store original user input in DB in English already (per your pipeline). If not, adapt accordingly.
            displayed_text = translate(text_report_en, TARGET_LANG)
            st.text_area(translate("Text Report", TARGET_LANG), displayed_text, height=140)

            # Current status (translate status)
            st.write(translate(f"Current Status: {current_status_en}", TARGET_LANG))

            # Analysis tools header
            st.markdown("### " + translate("Analysis Tools", TARGET_LANG))

            col_video, col_text = st.columns(2)

            # -------------------------
            # Video analysis (kept English)
            # -------------------------
            with col_video:
                analyze_video_label = translate(f"Analyze Video Report {report_id}", TARGET_LANG)
                # Provide a unique key per report for the button
                if st.button(analyze_video_label, key=f"analyze_video_btn_{report_id}"):
                    with st.spinner(translate("Analyzing video...", TARGET_LANG)):
                        try:
                            result = analyze_video(video_path)
                            st.session_state[f"video_analysis_{report_id}"] = result
                            st.success(translate("Video analysis complete!", TARGET_LANG))
                        except Exception as e:
                            st.error(translate(f"Error analyzing video: {str(e)}", TARGET_LANG))

                # Show video analysis results (ENGLISH; do not translate per choice)
                if st.session_state.get(f"video_analysis_{report_id}"):
                    st.markdown("" + translate("Video Analysis Results:", TARGET_LANG) + "")
                    st.write(st.session_state[f"video_analysis_{report_id}"])  # keep as-is (English)

            # -------------------------
            # Text analysis via Gemini
            # -------------------------
            with col_text:
                analyze_text_label = translate(f"Analyze Text Report {report_id}", TARGET_LANG)
                if st.button(analyze_text_label, key=f"analyze_text_btn_{report_id}"):
                    with st.spinner(translate("Analyzing text...", TARGET_LANG)):
                        if gemini_llm:
                            try:
                                # Process with Gemini (expected to return English structured output)
                                gemini_output_en = process_question_with_doc(text_report_en, gemini_llm)
                                st.session_state[f"text_analysis_en_{report_id}"] = gemini_output_en
                                st.success(translate("Text analysis complete!", TARGET_LANG))
                            except Exception as e:
                                st.error(translate(f"Error analyzing text: {str(e)}", TARGET_LANG))
                        else:
                            st.error(translate("Gemini LLM not initialized", TARGET_LANG))

                # Display text analysis results: TRANSLATE Gemini output into TARGET_LANG for UI
                if st.session_state.get(f"text_analysis_en_{report_id}"):
                    st.markdown("" + translate("Text Analysis Results:", TARGET_LANG) + "")
                    gemini_en = st.session_state[f"text_analysis_en_{report_id}"]
                    # Translate Gemini output for display
                    gemini_translated = translate(str(gemini_en), TARGET_LANG)
                    st.write(gemini_translated)

            # -------------------------
            # Accept / Reject
            # -------------------------
            st.markdown("### " + translate("Make Decision", TARGET_LANG))
            col1, col2 = st.columns(2)

            # Accept
            with col1:
                accept_label = translate(f"Accept Report {report_id}", TARGET_LANG)
                if st.button(accept_label, key=f"accept_btn_{report_id}"):
                    with st.spinner(translate("Processing acceptance...", TARGET_LANG)):
                        update_status(report_id, "Accepted")
                        updated = verify_status(report_id)
                        if updated == "Accepted":
                            # Generate docx: pass English gemini output and English status; generator will translate into TARGET_LANG when saving
                            file_path = generate_report(
                                report_id,
                                st.session_state.get(f"video_analysis_{report_id}"),
                                st.session_state.get(f"text_analysis_en_{report_id}"),
                                "Accepted",
                            )
                            st.success(translate(f"Report {report_id} has been accepted.", TARGET_LANG))
                            # Provide download button (label translated)
                            with open(file_path, "rb") as f:
                                st.download_button(
                                    label=translate("Download Accepted Report", TARGET_LANG),
                                    data=f,
                                    file_name=f"accepted_report_{report_id}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_accepted_{report_id}",
                                )
                        else:
                            st.error(translate("Status update failed. Please try again.", TARGET_LANG))

            # Reject
            with col2:
                reject_label = translate(f"Reject Report {report_id}", TARGET_LANG)
                if st.button(reject_label, key=f"reject_btn_{report_id}"):
                    with st.spinner(translate("Processing rejection...", TARGET_LANG)):
                        update_status(report_id, "Rejected")
                        updated = verify_status(report_id)
                        if updated == "Rejected":
                            file_path = generate_report(
                                report_id,
                                st.session_state.get(f"video_analysis_{report_id}"),
                                st.session_state.get(f"text_analysis_en_{report_id}"),
                                "Rejected",
                            )
                            st.error(translate(f"Report {report_id} has been rejected.", TARGET_LANG))
                            with open(file_path, "rb") as f:
                                st.download_button(
                                    label=translate("Download Rejected Report", TARGET_LANG),
                                    data=f,
                                    file_name=f"rejected_report_{report_id}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_rejected_{report_id}",
                                )
                        else:
                            st.error(translate("Status update failed. Please try again.", TARGET_LANG))

            st.divider()

# -------------------------
# Cleanup on exit
# -------------------------
def cleanup():
    try:
        conn.close()
    except Exception:
        pass


atexit.register(cleanup)